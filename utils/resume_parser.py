import streamlit as st
import pdfplumber
from docx import Document
import io
import re
from typing import Dict, Optional, List
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class ResumeParser:
    """Enhanced resume parser with multiple format support"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']

    def parse_file(self, uploaded_file) -> Dict[str, str]:
        """Parse uploaded file and extract text content"""
        if uploaded_file is None:
            return {"error": "No file uploaded"}

        file_extension = Path(uploaded_file.name).suffix.lower()

        try:
            if file_extension == '.pdf':
                return self._parse_pdf(uploaded_file)
            elif file_extension == '.docx':
                return self._parse_docx(uploaded_file)
            elif file_extension == '.txt':
                return self._parse_txt(uploaded_file)
            else:
                return {"error": f"Unsupported file format: {file_extension}"}

        except Exception as e:
            logger.error(f"Error parsing file {uploaded_file.name}: {e}")
            return {"error": f"Failed to parse file: {str(e)}"}

    def _parse_pdf(self, uploaded_file) -> Dict[str, str]:
        """Extract text from PDF file"""
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            text_content = ""

            with pdfplumber.open(uploaded_file) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num} ---\n"
                        text_content += page_text
                        text_content += "\n"

            if not text_content.strip():
                return {"error": "No text could be extracted from PDF"}

            return {
                "raw_text": text_content,
                "file_type": "pdf",
                "page_count": len(pdf.pages),
                "file_name": uploaded_file.name
            }

        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            return {"error": f"PDF parsing failed: {str(e)}"}

    def _parse_docx(self, uploaded_file) -> Dict[str, str]:
        """Extract text from DOCX file"""
        try:
            uploaded_file.seek(0)
            doc = Document(uploaded_file)

            text_content = ""
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1

            # Extract text from tables if any
            table_content = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_content += " | ".join(row_text) + "\n"

            if table_content:
                text_content += "\n--- Table Content ---\n" + table_content

            if not text_content.strip():
                return {"error": "No text could be extracted from DOCX"}

            return {
                "raw_text": text_content,
                "file_type": "docx",
                "paragraph_count": paragraph_count,
                "table_count": len(doc.tables),
                "file_name": uploaded_file.name
            }

        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return {"error": f"DOCX parsing failed: {str(e)}"}

    def _parse_txt(self, uploaded_file) -> Dict[str, str]:
        """Extract text from TXT file"""
        try:
            uploaded_file.seek(0)

            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text_content = None

            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    text_content = uploaded_file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                return {"error": "Could not decode text file with any supported encoding"}

            if not text_content.strip():
                return {"error": "Text file appears to be empty"}

            line_count = len(text_content.split('\n'))

            return {
                "raw_text": text_content,
                "file_type": "txt",
                "line_count": line_count,
                "file_name": uploaded_file.name
            }

        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            return {"error": f"TXT parsing failed: {str(e)}"}

    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume text"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "location": None
        }

        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()

        # Phone extraction (various formats)
        phone_patterns = [
            r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b',
            r'\b[0-9]{3}-[0-9]{3}-[0-9]{4}\b',
            r'\b\([0-9]{3}\) [0-9]{3}-[0-9]{4}\b'
        ]

        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info["phone"] = phone_match.group()
                break

        # LinkedIn extraction
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([A-Za-z0-9-]+)'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info["linkedin"] = f"linkedin.com/in/{linkedin_match.group(1)}"

        # Location extraction (basic city, state pattern)
        location_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\b'
        location_match = re.search(location_pattern, text)
        if location_match:
            contact_info["location"] = location_match.group()

        return contact_info

    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract common resume sections"""
        sections = {}

        # Common section headers - using raw strings
        section_patterns = {
            "summary": r'(?:professional\s+summary|summary|objective|profile)\s*:?\s*\n(.*?)(?=\n\s*(?:[A-Z][A-Z\s]+|$))',
            "experience": r'(?:work\s+experience|experience|employment|professional\s+experience)\s*:?\s*\n(.*?)(?=\n\s*(?:[A-Z][A-Z\s]+|$))',
            "education": r'(?:education|academic\s+background)\s*:?\s*\n(.*?)(?=\n\s*(?:[A-Z][A-Z\s]+|$))',
            "skills": r'(?:skills|technical\s+skills|core\s+competencies)\s*:?\s*\n(.*?)(?=\n\s*(?:[A-Z][A-Z\s]+|$))',
            "certifications": r'(?:certifications|certificates|licenses)\s*:?\s*\n(.*?)(?=\n\s*(?:[A-Z][A-Z\s]+|$))'
        }

        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(1).strip()

        return sections

    def get_file_stats(self, parsed_data: Dict) -> Dict[str, any]:
        """Get file statistics for analysis"""
        if "error" in parsed_data:
            return {"error": parsed_data["error"]}

        text = parsed_data.get("raw_text", "")

        stats = {
            "character_count": len(text),
            "word_count": len(text.split()),
            "line_count": len(text.split('\n')),
            "file_type": parsed_data.get("file_type", "unknown"),
            "file_name": parsed_data.get("file_name", "unknown")
        }

        # Add format-specific stats
        if parsed_data.get("file_type") == "pdf":
            stats["page_count"] = parsed_data.get("page_count", 0)
        elif parsed_data.get("file_type") == "docx":
            stats["paragraph_count"] = parsed_data.get("paragraph_count", 0)
            stats["table_count"] = parsed_data.get("table_count", 0)

        return stats

    def validate_resume_content(self, text: str) -> Dict[str, any]:
        """Validate resume content for completeness"""
        validation = {
            "has_contact_info": False,
            "has_experience": False,
            "has_education": False,
            "has_skills": False,
            "word_count_ok": False,
            "issues": []
        }

        # Check for contact information
        contact_info = self.extract_contact_info(text)
        if contact_info["email"] or contact_info["phone"]:
            validation["has_contact_info"] = True
        else:
            validation["issues"].append("Missing contact information (email or phone)")

        # Check for key sections
        sections = self.extract_sections(text)

        if sections.get("experience"):
            validation["has_experience"] = True
        else:
            validation["issues"].append("Missing work experience section")

        if sections.get("education"):
            validation["has_education"] = True
        else:
            validation["issues"].append("Missing education section")

        if sections.get("skills"):
            validation["has_skills"] = True
        else:
            validation["issues"].append("Missing skills section")

        # Check word count
        word_count = len(text.split())
        if 200 <= word_count <= 1000:
            validation["word_count_ok"] = True
        else:
            if word_count < 200:
                validation["issues"].append("Resume too short (< 200 words)")
            else:
                validation["issues"].append("Resume too long (> 1000 words)")

        return validation

# Global parser instance
resume_parser = ResumeParser()
